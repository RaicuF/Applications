import os
import json
import csv
from datetime import datetime
from typing import Dict, Any, List
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportService:
    def __init__(self):
        self.export_dir = "/tmp/exports"
        os.makedirs(self.export_dir, exist_ok=True)
    
    async def export_pdf(self, data: Dict[str, Any]) -> str:
        filename = f"{self.export_dir}/analysis_{data['id']}.pdf"
        doc = SimpleDocTemplate(filename, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a2e'),
            alignment=TA_CENTER
        )
        
        story.append(Paragraph("Log Analysis Report", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        info_data = [
            ["Report ID:", data['id']],
            ["Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["Total Logs:", str(data['total_logs'])],
            ["Risk Level:", data.get('ai_insights', {}).get('summary', {}).get('risk_level', 'N/A')]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))
        
        if data.get('suspicious_ips'):
            story.append(Paragraph("Suspicious IPs Detected", styles['Heading2']))
            ip_data = [["IP Address", "Risk Score", "Source"]]
            for ip in data['suspicious_ips'][:10]:
                ip_data.append([
                    ip.get('ip', ''),
                    str(ip.get('risk_score', 0)),
                    ', '.join(ip.get('sources', {}).keys())
                ])
            
            ip_table = Table(ip_data, colWidths=[2*inch, 1.5*inch, 2.5*inch])
            ip_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            story.append(ip_table)
            story.append(Spacer(1, 0.3*inch))
        
        if data.get('ai_insights'):
            insights = data['ai_insights']
            story.append(Paragraph("AI Analysis Insights", styles['Heading2']))
            
            if insights.get('concerns'):
                story.append(Paragraph("Security Concerns:", styles['Heading3']))
                for concern in insights['concerns']:
                    story.append(Paragraph(f"• {concern}", styles['Normal']))
            
            if insights.get('recommendations'):
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph("Recommendations:", styles['Heading3']))
                for rec in insights['recommendations']:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))
        
        doc.build(story)
        return filename
    
    async def export_excel(self, data: Dict[str, Any]) -> str:
        filename = f"{self.export_dir}/analysis_{data['id']}.xlsx"
        
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            summary_df = pd.DataFrame([{
                'Report ID': data['id'],
                'Generated': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Total Logs': data['total_logs'],
                'Risk Level': data.get('ai_insights', {}).get('summary', {}).get('risk_level', 'N/A')
            }])
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            if data.get('logs'):
                logs_df = pd.DataFrame(data['logs'][:10000])
                logs_df.to_excel(writer, sheet_name='Logs', index=False)
            
            if data.get('suspicious_ips'):
                ips_df = pd.DataFrame(data['suspicious_ips'])
                ips_df.to_excel(writer, sheet_name='Suspicious IPs', index=False)
            
            if data.get('ai_insights', {}).get('anomalies'):
                anomalies_df = pd.DataFrame(data['ai_insights']['anomalies'])
                anomalies_df.to_excel(writer, sheet_name='Anomalies', index=False)
            
            if data.get('ai_insights', {}).get('statistics'):
                stats_df = pd.DataFrame([data['ai_insights']['statistics']])
                stats_df.to_excel(writer, sheet_name='Statistics', index=False)
        
        return filename
    
    async def export_word(self, data: Dict[str, Any]) -> str:
        filename = f"{self.export_dir}/analysis_{data['id']}.docx"
        doc = Document()
        
        title = doc.add_heading('Log Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Report Information', level=1)
        doc.add_paragraph(f"Report ID: {data['id']}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Logs Analyzed: {data['total_logs']}")
        
        risk_level = data.get('ai_insights', {}).get('summary', {}).get('risk_level', 'N/A')
        p = doc.add_paragraph()
        p.add_run('Risk Level: ').bold = True
        run = p.add_run(risk_level.upper())
        if risk_level == 'critical':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif risk_level == 'high':
            run.font.color.rgb = RGBColor(255, 165, 0)
        elif risk_level == 'medium':
            run.font.color.rgb = RGBColor(255, 255, 0)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)
        
        if data.get('suspicious_ips'):
            doc.add_heading('Suspicious IP Addresses', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'IP Address'
            hdr_cells[1].text = 'Risk Score'
            hdr_cells[2].text = 'Source'
            
            for ip in data['suspicious_ips'][:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = ip.get('ip', '')
                row_cells[1].text = str(ip.get('risk_score', 0))
                row_cells[2].text = ', '.join(ip.get('sources', {}).keys())
        
        if data.get('ai_insights'):
            doc.add_heading('AI Analysis Insights', level=1)
            insights = data['ai_insights']
            
            if insights.get('concerns'):
                doc.add_heading('Security Concerns', level=2)
                for concern in insights['concerns']:
                    doc.add_paragraph(concern, style='List Bullet')
            
            if insights.get('patterns'):
                doc.add_heading('Detected Patterns', level=2)
                for pattern in insights['patterns']:
                    if isinstance(pattern, dict):
                        doc.add_paragraph(f"• {pattern.get('description', str(pattern))}", style='List Bullet')
                    else:
                        doc.add_paragraph(f"• {pattern}", style='List Bullet')
            
            if insights.get('recommendations'):
                doc.add_heading('Recommendations', level=2)
                for i, rec in enumerate(insights['recommendations'], 1):
                    doc.add_paragraph(f"{i}. {rec}")
        
        doc.save(filename)
        return filename
    
    async def export_csv(self, data: Dict[str, Any]) -> str:
        filename = f"{self.export_dir}/analysis_{data['id']}.csv"
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            if data.get('logs'):
                if data['logs']:
                    fieldnames = list(data['logs'][0].keys())
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    for log in data['logs']:
                        writer.writerow(log)
            else:
                writer = csv.writer(csvfile)
                writer.writerow(['Report ID', 'Generated', 'Total Logs', 'Risk Level'])
                writer.writerow([
                    data['id'],
                    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    data['total_logs'],
                    data.get('ai_insights', {}).get('summary', {}).get('risk_level', 'N/A')
                ])
                
                if data.get('suspicious_ips'):
                    writer.writerow([])
                    writer.writerow(['Suspicious IPs'])
                    writer.writerow(['IP Address', 'Risk Score', 'Is Malicious'])
                    for ip in data['suspicious_ips']:
                        writer.writerow([
                            ip.get('ip', ''),
                            ip.get('risk_score', 0),
                            ip.get('is_malicious', False)
                        ])
        
        return filename
    
    async def export_anomalies_csv(self, data: Dict[str, Any]) -> str:
        filename = f"{self.export_dir}/anomalies_{data['id']}.csv"
        
        anomalies = data.get('ai_insights', {}).get('anomalies', [])
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            if anomalies:
                fieldnames = ['type', 'severity', 'description', 'details']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                
                for anomaly in anomalies:
                    writer.writerow({
                        'type': anomaly.get('type', ''),
                        'severity': anomaly.get('severity', ''),
                        'description': anomaly.get('description', ''),
                        'details': json.dumps(anomaly.get('details', {}))
                    })
            else:
                writer = csv.writer(csvfile)
                writer.writerow(['No anomalies detected'])
        
        return filename